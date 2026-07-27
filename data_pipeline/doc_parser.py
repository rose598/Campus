"""
doc_parser.py — 文档解析器（PDF / Word / HTML）

职责:
  - 根据文件扩展名自动选择解析策略
  - PDF: pypdfium2 提取文本（保留页码信息）
  - Word (.docx): python-docx 提取段落文本
  - HTML: BeautifulSoup 提取正文，过滤导航/脚本/样式
  - 统一输出 ParsedDocument，与下游 CampusDocument 对齐

使用方式:
  parser = DocParser()
  doc = parser.parse("path/to/file.pdf")
  print(doc.title, doc.content[:100])
"""

from __future__ import annotations

import hashlib
import logging
import re
from pathlib import Path
from typing import Dict, List, Optional

from pydantic import BaseModel, Field

logger = logging.getLogger(__name__)


# ─────────────────────────────────────────────
#  中间数据结构
# ─────────────────────────────────────────────

class ParsedDocument(BaseModel):
    """解析后的文档结构，供下游 TextCleaner / Chunker / Annotator 消费"""
    doc_id: str = Field(..., description="文档唯一 ID，格式 DOC_XXXXXXXX")
    title: str = Field(default="", description="文档标题（从文件名或元数据提取）")
    content: str = Field(..., min_length=1, description="提取的原始文本")
    source_path: str = Field(..., description="文件本地路径")
    source_url: Optional[str] = Field(None, description="来源 URL（HTML 页面）")
    file_type: str = Field(..., description="文件类型: pdf / docx / html")
    page_count: int = Field(default=0, ge=0, description="页数（PDF 专用，其他格式为 0）")
    metadata: Dict = Field(default_factory=dict, description="额外元数据（作者/创建时间等）")


# ─────────────────────────────────────────────
#  文档解析器
# ─────────────────────────────────────────────

class DocParser:
    """
    统一文档解析接口。

    支持格式：
      - .pdf  （pypdfium2）
      - .docx （python-docx）
      - .html / .htm （BeautifulSoup）

    示例:
        >>> parser = DocParser()
        >>> doc = parser.parse("data/raw/policies/保研政策.pdf")
        >>> doc.file_type
        'pdf'
    """

    # 支持的文件扩展名 → 解析方法映射
    _SUPPORTED_EXTS = {
        ".pdf": "_parse_pdf",
        ".docx": "_parse_docx",
        ".html": "_parse_html",
        ".htm": "_parse_html",
    }

    # ── 公共接口 ──────────────────────────────

    def parse(self, file_path: str | Path) -> ParsedDocument:
        """
        解析单个文件，返回 ParsedDocument。

        Args:
            file_path: 文件路径（str 或 Path）

        Raises:
            FileNotFoundError: 文件不存在
            ValueError: 不支持的文件格式
            RuntimeError: 解析过程出错
        """
        path = Path(file_path)
        if not path.exists():
            raise FileNotFoundError(f"文件不存在: {path}")

        ext = path.suffix.lower()
        if ext not in self._SUPPORTED_EXTS:
            raise ValueError(
                f"不支持的文件格式 '{ext}'，当前支持: {list(self._SUPPORTED_EXTS.keys())}"
            )

        method_name = self._SUPPORTED_EXTS[ext]
        method = getattr(self, method_name)

        logger.info("[DocParser] 开始解析: %s (类型=%s)", path.name, ext)
        try:
            result: ParsedDocument = method(path)
            logger.info(
                "[DocParser] 解析完成: %s | 页数=%d | 内容长度=%d",
                result.title, result.page_count, len(result.content),
            )
            return result
        except Exception as exc:
            logger.error("[DocParser] 解析失败 [%s]: %s", path.name, exc, exc_info=True)
            raise RuntimeError(f"解析文件失败 '{path}': {exc}") from exc

    def parse_batch(self, file_paths: List[str | Path]) -> List[ParsedDocument]:
        """
        批量解析，跳过失败的文件（不中断整体流程）。

        Returns:
            成功解析的 ParsedDocument 列表
        """
        results: List[ParsedDocument] = []
        for fp in file_paths:
            try:
                results.append(self.parse(fp))
            except (RuntimeError, FileNotFoundError, ValueError) as exc:
                logger.warning("[DocParser] 跳过文件 %s: %s", fp, exc)
        logger.info("[DocParser] 批量解析完成: %d/%d 成功", len(results), len(file_paths))
        return results

    # ── PDF 解析 ──────────────────────────────

    def _parse_pdf(self, path: Path) -> ParsedDocument:
        """使用 pypdfium2 提取 PDF 文本"""
        import pypdfium2 as pdfium

        pdf = pdfium.PdfDocument(str(path))
        page_count = len(pdf)
        pages_text: List[str] = []

        for i, page in enumerate(pdf):
            textpage = page.get_textpage()
            text = textpage.get_text_range()
            textpage.close()
            page.close()
            if text and text.strip():
                pages_text.append(text)

        pdf.close()

        full_text = "\n\n".join(pages_text)
        if not full_text.strip():
            logger.warning("[DocParser] PDF 文本为空: %s", path.name)
            full_text = ""  # TextCleaner 会处理空文档兜底

        # 从 PDF 元数据提取标题
        title = self._extract_pdf_title(pdf, path)

        metadata: Dict = {"page_count": page_count}
        # pypdfium2 元数据访问
        try:
            meta = pdf.get_metadata_dict()
            if meta:
                metadata["pdf_metadata"] = {
                    k: str(v) for k, v in meta.items() if v
                }
        except Exception:
            pass

        return ParsedDocument(
            doc_id=self._generate_doc_id(path),
            title=title,
            content=full_text,
            source_path=str(path),
            file_type="pdf",
            page_count=page_count,
            metadata=metadata,
        )

    @staticmethod
    def _extract_pdf_title(pdf, path: Path) -> str:
        """从 PDF 元数据或文件名提取标题"""
        try:
            meta = pdf.get_metadata_dict()
            title = meta.get("Title") or meta.get("title")
            if title and str(title).strip():
                return str(title).strip()
        except Exception:
            pass
        # 回退到文件名（去扩展名）
        return path.stem

    # ── Word (.docx) 解析 ────────────────────

    def _parse_docx(self, path: Path) -> ParsedDocument:
        """使用 python-docx 提取 Word 文档文本"""
        from docx import Document

        doc = Document(str(path))
        paragraphs: List[str] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                paragraphs.append(text)

        # 也提取表格内容（政策文件常以表格形式呈现）
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        full_text = "\n\n".join(paragraphs)
        if not full_text.strip():
            logger.warning("[DocParser] DOCX 文本为空: %s", path.name)
            full_text = ""

        # 标题：优先取 core_properties.title，回退到文件名
        title = path.stem
        try:
            props = doc.core_properties
            if props.title and props.title.strip():
                title = props.title.strip()
        except Exception:
            pass

        metadata: Dict = {}
        try:
            props = doc.core_properties
            metadata.update({
                k: str(v)
                for k, v in {
                    "author": props.author,
                    "created": props.created,
                    "modified": props.modified,
                    "subject": props.subject,
                }.items()
                if v
            })
        except Exception:
            pass

        return ParsedDocument(
            doc_id=self._generate_doc_id(path),
            title=title,
            content=full_text,
            source_path=str(path),
            file_type="docx",
            page_count=0,
            metadata=metadata,
        )

    # ── HTML 解析 ─────────────────────────────

    def _parse_html(self, path: Path) -> ParsedDocument:
        """使用 BeautifulSoup 提取 HTML 正文，过滤脚本/样式/导航"""
        from bs4 import BeautifulSoup

        with open(path, "r", encoding="utf-8", errors="replace") as f:
            html = f.read()

        soup = BeautifulSoup(html, "html.parser")

        # 移除脚本、样式、导航、页脚等非正文元素
        for tag in soup.find_all(["script", "style", "nav", "footer", "header", "aside", "iframe"]):
            tag.decompose()

        # 提取标题
        title = path.stem
        title_tag = soup.find("title")
        if title_tag and title_tag.get_text(strip=True):
            title = title_tag.get_text(strip=True)
        h1_tag = soup.find("h1")
        if h1_tag and h1_tag.get_text(strip=True):
            title = h1_tag.get_text(strip=True)

        # 提取正文文本 —— 逐块级元素提取，用双换行拼接，便于下游分段清洗
        body = soup.find("body")
        target = body if body else soup
        _BLOCK_TAGS = {"p", "div", "h1", "h2", "h3", "h4", "h5", "h6",
                       "li", "tr", "blockquote", "pre", "td"}
        blocks = target.find_all(_BLOCK_TAGS)
        if blocks:
            paragraphs = [b.get_text(strip=True) for b in blocks]
            full_text = "\n\n".join(p for p in paragraphs if p)
        else:
            # 无块级标签时回退到整体提取
            full_text = target.get_text(separator="\n", strip=True)

        if not full_text.strip():
            logger.warning("[DocParser] HTML 文本为空: %s", path.name)
            full_text = ""

        # 提取 meta 信息
        metadata: Dict = {}
        meta_desc = soup.find("meta", attrs={"name": "description"})
        if meta_desc and meta_desc.get("content"):
            metadata["description"] = meta_desc["content"]
        meta_kw = soup.find("meta", attrs={"name": "keywords"})
        if meta_kw and meta_kw.get("content"):
            metadata["keywords"] = meta_kw["content"]

        # 尝试从 HTML 中提取 source_url（canonical link）
        source_url: Optional[str] = None
        canonical = soup.find("link", attrs={"rel": "canonical"})
        if canonical and canonical.get("href"):
            source_url = str(canonical["href"])

        return ParsedDocument(
            doc_id=self._generate_doc_id(path),
            title=title,
            content=full_text,
            source_path=str(path),
            source_url=source_url,
            file_type="html",
            page_count=0,
            metadata=metadata,
        )

    # ── 工具方法 ──────────────────────────────

    @staticmethod
    def _generate_doc_id(path: Path) -> str:
        """
        基于文件路径生成稳定的 doc_id（格式: DOC_XXXXXXXX）。
        使用路径哈希的前 8 位十六进制。
        """
        digest = hashlib.md5(str(path.resolve()).encode("utf-8")).hexdigest()[:8]
        return f"DOC_{digest.upper()}"


# ─────────────────────────────────────────────
#  命令行快速测试入口
# ─────────────────────────────────────────────

if __name__ == "__main__":
    import sys

    logging.basicConfig(level=logging.INFO, format="%(levelname)s - %(message)s")

    if len(sys.argv) < 2:
        print("用法: python doc_parser.py <file_path> [file_path2 ...]")
        sys.exit(1)

    parser = DocParser()
    for fp in sys.argv[1:]:
        try:
            doc = parser.parse(fp)
            print(f"✓ {doc.title} [{doc.file_type}] | "
                  f"{len(doc.content)} chars | {doc.page_count} pages")
            print(f"  doc_id: {doc.doc_id}")
            print(f"  预览: {doc.content[:200]}...")
            print()
        except Exception as e:
            print(f"✗ {fp}: {e}")
